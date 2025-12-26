from __future__ import annotations

import io
import logging
import re
from typing import Final

from docx import Document
from pypdf import PdfReader

from app.core.logging import get_logger

logger = get_logger(__name__)

# Разрешённые / ожидаемые content-type
DOCX_MIME_TYPES: Final[set[str]] = {
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "application/msword",
}
PDF_MIME_TYPES: Final[set[str]] = {
    "application/pdf",
}
TEXT_MIME_TYPES: Final[set[str]] = {
    "text/plain",
    "text/markdown",
    "text/csv",
    "text/html",  # только если вы уверены, что хотите HTML как сырой текст
}


def _normalize_text(text: str) -> str:
    """
    Нормализует текст перед сохранением в БД:
    - удаляет NUL (\x00), которые Postgres не принимает
    - убирает управляющие символы кроме табуляции и перевода строки
    - нормализует переводы строк и лишние пробелы
    """
    if not text:
        return ""

    # Удаляем NUL-символы
    text = text.replace("\x00", "")

    # Удаляем прочие непечатаемые управляющие символы (кроме \n, \r, \t)
    control_chars_regex = r"[\u0000-\u0008\u000B\u000C\u000E-\u001F\u007F]"
    text = re.sub(control_chars_regex, "", text)

    # Нормализуем переводы строк (CRLF -> LF)
    text = text.replace("\r\n", "\n").replace("\r", "\n")

    # Сжимаем более трёх переводов строки подряд до двух
    text = re.sub(r"\n{3,}", "\n\n", text)

    # Удаляем лишние пробелы в начале и в конце строк
    lines = [line.strip() for line in text.split("\n")]
    text = "\n".join(lines).strip()

    return text


def extract_text_from_docx(file_bytes: bytes, filename: str | None = None) -> str:
    """
    Извлекает текст из DOCX файла с помощью python-docx.
    """
    logger.info("Extracting text from DOCX file", extra={"source_filename": filename})

    with io.BytesIO(file_bytes) as buffer:
        document = Document(buffer)

    paragraphs: list[str] = []

    # Основные абзацы
    for para in document.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)

    # Таблицы (если в документе есть важный текст в таблицах)
    for table in document.tables:
        for row in table.rows:
            cells_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells_text:
                paragraphs.append(" | ".join(cells_text))

    raw_text = "\n".join(paragraphs)
    normalized = _normalize_text(raw_text)

    logger.info(
        "DOCX text extracted",
        extra={
            "source_filename": filename,
            "chars_raw": len(raw_text),
            "chars_normalized": len(normalized),
        },
    )

    return normalized


def extract_text_from_pdf(file_bytes: bytes, filename: str | None = None) -> str:
    """
    Извлекает текст из PDF с помощью pypdf.
    """
    logger.info("Extracting text from PDF file", extra={"source_filename": filename})

    with io.BytesIO(file_bytes) as buffer:
        reader = PdfReader(buffer)

        pages_text: list[str] = []
        for page_index, page in enumerate(reader.pages):
            try:
                page_text = page.extract_text() or ""
            except Exception as exc:
                logger.warning(
                    "Failed to extract text from PDF page",
                    extra={
                        "source_filename": filename,
                        "page_index": page_index,
                        "error": str(exc),
                    },
                )
                page_text = ""
            page_text = page_text.strip()
            if page_text:
                pages_text.append(page_text)

    raw_text = "\n\n".join(pages_text)
    normalized = _normalize_text(raw_text)

    logger.info(
        "PDF text extracted",
        extra={
            "source_filename": filename,
            "pages": len(reader.pages),
            "chars_raw": len(raw_text),
            "chars_normalized": len(normalized),
        },
    )

    return normalized


def extract_text_from_plain(
    file_bytes: bytes,
    filename: str | None = None,
    encoding: str = "utf-8",
) -> str:
    """
    Извлекает текст из простого текстового файла.
    """
    logger.info(
        "Extracting text from plain text file",
        extra={"source_filename": filename, "encoding": encoding},
    )

    try:
        text = file_bytes.decode(encoding)
    except Exception as exc:
        logger.warning(
            "Failed to decode text file with encoding, using errors='ignore'",
            extra={"source_filename": filename, "encoding": encoding, "error": str(exc)},
        )
        text = file_bytes.decode(encoding, errors="ignore")

    normalized = _normalize_text(text)

    logger.info(
        "Plain text extracted",
        extra={
            "source_filename": filename,
            "chars_raw": len(text),
            "chars_normalized": len(normalized),
        },
    )

    return normalized


def extract_text(
    file_bytes: bytes, content_type: str | None, filename: str | None
) -> str:
    """
    Высокоуровневая функция определения формата и извлечения текста.
    """
    content_type = (content_type or "").lower()
    filename = (filename or "").lower()

    logger.info(
        "Starting text extraction",
        extra={"source_filename": filename, "content_type": content_type},
    )

    try:
        # DOCX
        if content_type in DOCX_MIME_TYPES or (
            filename and (filename.endswith(".docx") or filename.endswith(".doc"))
        ):
            return extract_text_from_docx(file_bytes, filename=filename)

        # PDF
        if content_type in PDF_MIME_TYPES or (filename and filename.endswith(".pdf")):
            return extract_text_from_pdf(file_bytes, filename=filename)

        # TXT и подобные
        if content_type in TEXT_MIME_TYPES or (filename and filename.endswith(".txt")):
            return extract_text_from_plain(file_bytes, filename=filename)

        # Fallback: пробуем как текстовый файл
        logger.info(
            "Falling back to plain text extractor for unknown content type",
            extra={"source_filename": filename, "content_type": content_type},
        )
        return extract_text_from_plain(file_bytes, filename=filename)
    except Exception as exc:
        # В проде важно не уронить весь процесс: логируем и возвращаем пустую строку
        logger.error(
            "Failed to extract text from file",
            extra={
                "source_filename": filename,
                "content_type": content_type,
                "error": str(exc),
            },
        )
        return ""

